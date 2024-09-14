import json
import os
import tempfile
from enum import StrEnum
from time import sleep

import lancedb
import pandas as pd
from burr.core import Application
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from fasthtml.common import (
    H1,
    A,
    Button,
    Container,
    Div,
    Form,
    Html,
    Input,
    Li,
    Ol,
    Option,
    P,
    Script,
    Select,
    Style,
    Ul,
    fast_app,
    serve,
)
from loguru import logger
from lxml import etree
from starlette.background import BackgroundTask
from starlette.requests import Request
from starlette.responses import FileResponse

from dreamai.rag import application
from dreamai.rag_utils import add_data_with_descriptions
from dreamai.settings import ModalSettings
from dreamai.utils import flatten_list

modal_settings = ModalSettings()

SOURCE_TEXT_SIZE = 100_000
LANCE_DIR = modal_settings.lance_dir
MODEL = modal_settings.model


logger.info(f"LANCE_DIR: {LANCE_DIR}")
logger.info(f"MODEL: {MODEL}")


style = """
.spinner-container {
    margin-top: 7px;
    margin-bottom: 10px;
    display: none;
}
.spinner {
    border: 4px solid #f3f3f3;
    border-top: 4px solid #3498db;
    border-radius: 50%;
    width: 30px;
    height: 30px;
    animation: spin 1s linear infinite;
    display: inline-block;
    margin-right: 10px;
    vertical-align: middle;
}
@keyframes spin {
    0% { transform: rotate(0deg); }
    100% { transform: rotate(360deg); }
}
.section {
    border: 1px solid #e0e0e0;
    border-radius: 8px;
    padding: 20px;
    margin-bottom: 20px;
    background-color: #f9f9f9;
}
.section-title {
    font-size: 18px;
    font-weight: bold;
    margin-bottom: 15px;
    color: #333;
}

.divider {
    border-top: 1px solid #e0e0e0;
    margin: 30px 0;
}
"""

table_descriptions = []
questions = []
qna = {}
app = fast_app(live=True, hdrs=(Style(style), Html(data_theme="light")))[0]


class Mode(StrEnum):
    RFP = "rfp"
    SECURITY = "security"


def Spinner(message):
    return Div(Div(cls="spinner"), message, cls="spinner-container", id="spinner")


def get_sorted_indexes():
    if not os.path.exists(LANCE_DIR):
        os.makedirs(LANCE_DIR)
    indexes = [d for d in os.listdir(LANCE_DIR) if os.path.isdir(os.path.join(LANCE_DIR, d))]
    return sorted(
        indexes, key=lambda x: os.path.getmtime(os.path.join(LANCE_DIR, x)), reverse=True
    )


@app.get("/")
async def root(request: Request):
    return await home(request, Mode.RFP)


@app.get("/{mode}")
async def home(request: Request, mode: Mode = Mode.RFP):
    logger.info(f"Mode: {mode}")
    current_index = request.query_params.get("index")
    indexes = get_sorted_indexes()
    if not current_index and indexes:
        current_index = indexes[0]

    index_selection = (
        Select(
            *[
                Option(index, value=index, selected=(index == current_index))
                for index in indexes
            ],
            Option("Create New Index", value="new"),
            name="index",
            id="index-select",
        )
        if indexes
        else Input(
            type="text",
            name="new_index",
            id="new-index-input",
            placeholder="New Index Name",
        )
    )

    return Container(
        H1(f"{mode.upper() if len(mode) <= 3 else mode.title()} Questionnaire Tool"),
        Div(
            Form(
                index_selection,
                Input(
                    type="text",
                    name="new_index",
                    id="new-index-input",
                    placeholder="New Index Name",
                    style="display: none;" if indexes else "",
                )
                if indexes
                else None,
                Button("Select/Create Index", type="submit", id="index-button"),
                hx_post="/select-index",
                hx_target="#index-info",
                hx_swap="innerHTML",
            ),
            Div(id="index-info"),
            cls="section",
        ),
        Div(cls="divider"),
        Div(
            Div(
                P(
                    "Upload your company documents (PDF, DOCX, TXT, or MD files)",
                    cls="section-title",
                ),
                Form(
                    Input(
                        type="file",
                        name="documents",
                        multiple=True,
                        accept=".pdf,.docx,.txt,.md",
                        id="document-input",
                    ),
                    Button(
                        "Upload Documents",
                        type="submit",
                        id="upload-button",
                        style="display: none;",
                    ),
                    Spinner("Uploading documents..."),
                    hx_post="/upload",
                    hx_target="#file-list",
                    hx_swap="beforeend",
                    enctype="multipart/form-data",
                ),
                Div(id="file-list"),
                id="document-upload-section",
            ),
            cls="section",
        ),
        Div(cls="divider"),
        Div(
            Div(
                P("Upload your Questions CSV file", cls="section-title"),
                Form(
                    Input(type="file", name="questions", accept=".csv", id="questions-input"),
                    Button(
                        "Upload Questions",
                        type="submit",
                        id="questions-upload-button",
                        style="display: none;",
                    ),
                    hx_post="/upload-questions",
                    hx_target="#questions-info",
                    hx_swap="innerHTML",
                    enctype="multipart/form-data",
                ),
                Div(id="questions-info"),
                id="questions-upload-section",
            ),
            cls="section",
        ),
        Div(
            Button(
                "Answer Questions",
                id="answer-button",
                hx_post="/answer",
                hx_target="#answering-status",
                style="display: none;",
            ),
            Spinner("Answering Questions..."),
            Div(id="answering-status"),
            id="answering-section",
        ),
        Script("""
        var indexSelect = document.getElementById('index-select');
        var newIndexInput = document.getElementById('new-index-input');
        
        if (indexSelect) {
            indexSelect.addEventListener('change', function() {
                if (newIndexInput) {
                    newIndexInput.style.display = this.value === 'new' ? 'block' : 'none';
                }
            });
        }
        
        document.body.addEventListener('htmx:beforeRequest', function(evt) {
            if (evt.detail.pathInfo.requestPath === '/upload') {
                document.querySelector('#document-upload-section .spinner-container').style.display = 'block';
            } else if (evt.detail.pathInfo.requestPath === '/answer') {
                document.querySelector('#answering-section .spinner-container').style.display = 'block';
            }
        });

        document.body.addEventListener('htmx:afterRequest', function(evt) {
            if (evt.detail.successful) {
                if (evt.detail.pathInfo.requestPath === '/upload') {
                    document.getElementById('document-input').value = '';
                    document.getElementById('upload-button').style.display = 'none';
                    document.querySelector('#document-upload-section .spinner-container').style.display = 'none';
                } else if (evt.detail.pathInfo.requestPath === '/upload-questions') {
                    document.getElementById('questions-input').value = '';
                    document.getElementById('questions-upload-button').style.display = 'none';
                } else if (evt.detail.pathInfo.requestPath === '/answer') {
                    document.querySelector('#answering-section .spinner-container').style.display = 'none';
                }
                updateProcessButton();
            }
        });

        document.getElementById('document-input').addEventListener('change', function(evt) {
            var uploadButton = document.getElementById('upload-button');
            uploadButton.style.display = this.files.length > 0 ? 'inline-block' : 'none';
        });

        document.getElementById('questions-input').addEventListener('change', function(evt) {
            var uploadButton = document.getElementById('questions-upload-button');
            uploadButton.style.display = this.files.length > 0 ? 'inline-block' : 'none';
        });

        function updateProcessButton() {
            var questionsInfoEmpty = document.getElementById('questions-info').children.length === 0;
            var processButton = document.getElementById('answer-button');
            processButton.style.display = !questionsInfoEmpty ? 'inline-block' : 'none';
        }

        updateProcessButton();

        // document.body.addEventListener('click', function(evt) {
        //     if (evt.target.id === 'download-button') {
        //         setTimeout(function() {
        //             var currentIndex = document.getElementById('index-select').value;
        //             var indexInfo = document.getElementById('index-info').innerHTML;
        //             var currentPath = window.location.pathname;
        //             htmx.ajax('GET', currentPath + '?index=' + encodeURIComponent(currentIndex), {
        //                 target: 'body',
        //                 swap: 'innerHTML',
        //                 complete: function() {
        //                     document.getElementById('index-info').innerHTML = indexInfo;
        //                     document.getElementById('index-select').value = currentIndex;
        //                 }
        //             });
        //         }, 1000);  // Wait for 1 second to ensure the download has started
        //     }
        // });
        
        document.addEventListener('DOMContentLoaded', function() {
            var indexSelect = document.getElementById('index-select');
            if (indexSelect) {
                var currentIndex = indexSelect.value;
                htmx.ajax('GET', '/set-index?index=' + encodeURIComponent(currentIndex), {
                    target: '#index-info',
                    swap: 'innerHTML'
                });
            } else {
                htmx.ajax('GET', '/set-index', {
                    target: '#index-info',
                    swap: 'innerHTML'
                });
            }
        });
        """),
    )


@app.get("/set-index")
async def set_index(request):
    index = request.query_params.get("index")
    if not index:
        indexes = get_sorted_indexes()
        index = indexes[0] if indexes else None

    if index:
        os.makedirs(os.path.join(LANCE_DIR, index), exist_ok=True)
        global lance_db
        lance_db = lancedb.connect(os.path.join(LANCE_DIR, index))
        existing_tables = list(lance_db.table_names())
        return Div(
            P(f"Selected/Created index: {index}"),
            P(f"Existing tables: {len(existing_tables)}"),
            Ul(*[Li(table) for table in existing_tables[:5]]),
            P("..." if len(existing_tables) > 5 else ""),
        )
    else:
        return Div("No index available. Please create a new one.")


@app.post("/select-index")
async def select_index(request):
    form_data = await request.form()
    selected_index = form_data.get("index")
    new_index = form_data.get("new_index")

    if not selected_index or selected_index == "new":
        selected_index = new_index

    if not selected_index:
        return Div("No index selected or created.", _="on load wait 2s then remove me")

    os.makedirs(os.path.join(LANCE_DIR, selected_index), exist_ok=True)

    global lance_db
    lance_db = lancedb.connect(os.path.join(LANCE_DIR, selected_index))

    existing_tables = list(lance_db.table_names())

    return Div(
        P(f"Selected/Created index: {selected_index}"),
        P(f"Existing tables: {len(existing_tables)}"),
        Ul(*[Li(table) for table in existing_tables[:5]]),
        P("..." if len(existing_tables) > 5 else ""),
    )


@app.post("/upload")
async def upload(request):
    global table_descriptions
    form = await request.form()
    files = form.getlist("documents")
    if not files:
        return Li("No files were uploaded.", _="on load wait 2s then remove me")

    uploaded_files = []
    with tempfile.TemporaryDirectory() as temp_dir:
        for file in files:
            file_name = file.filename
            file_path = os.path.join(temp_dir, file_name)
            with open(file_path, "wb") as f:
                f.write(file.file.read())

            # Check if the file has already been added to the database
            if lance_db.table_names() and file_name in lance_db.table_names():
                uploaded_files.append(f"{file_name} (already exists)")
                continue

            uploaded_files.append(file_name)
            add_data_with_descriptions(
                model=MODEL,
                lance_db=lance_db,
                data=file_path,
                table_descriptions=table_descriptions,
            )

        logger.info(f"Processed {len(files)} files")

    return Ul(*[Li(file) for file in uploaded_files])


@app.post("/upload-questions")
async def upload_questions(request):
    global questions
    form = await request.form()
    questions_file = form.get("questions")
    if not questions_file:
        return Div("No questions file was uploaded.", _="on load wait 2s then remove me")

    with tempfile.TemporaryDirectory() as temp_dir:
        file_path = os.path.join(temp_dir, questions_file.filename)
        with open(file_path, "wb") as f:
            f.write(questions_file.file.read())

        df = pd.read_csv(file_path)
        questions = df.iloc[:, 0].tolist()
        question_count = len(questions)

        logger.info(
            f"Uploaded questions file: {questions_file.filename} with {question_count} questions"
        )
        logger.info(f"Stored {question_count} questions in global variable")

    return Div(
        P(f"Questions file '{questions_file.filename}' uploaded successfully."),
        P(f"Number of questions: {question_count}"),
        Ol(
            *[
                Li(question[:100] + "..." if len(question) > 100 else question)
                for question in questions[:5]
            ]
        ),
        P("..." if question_count > 5 else ""),
    )


async def answer_questions(app: Application, questions: list[str]):
    global qna
    qna = {"questions": questions, "answers": [], "sources": []}
    for query in qna["questions"]:
        inputs = {"query": query}
        logger.info(f"\nProcessing query: {query}")
        while True:
            step_result = app.step(inputs=inputs)
            if step_result is None:
                logger.error("Error: app.step() returned None")
                break
            action, result, _ = step_result
            logger.info(f"\nAction: {action.name}\n")
            logger.success(f"RESULT: {result}\n")
            if action.name == "terminate":
                break
            elif action.name in ["update_chat_history"]:
                qna["answers"].append(result["chat_history"][-1]["content"])
                qna["sources"].append(
                    [
                        {k: v for k, v in json.loads(d).items() if k != "index"}
                        for d in set(
                            [
                                json.dumps(s, sort_keys=True)
                                for s in flatten_list(app.state.get("source_docs", []))
                            ]
                        )
                    ]
                )
                break


@app.post("/answer")
async def answer(request):
    app = application(db=lance_db, model=MODEL, has_web=False, only_data=True)
    await answer_questions(app=app, questions=questions)
    logger.info(f"\n\nQNA: {qna}\n\n")
    return Div(
        Div(cls="divider"),
        A(
            "Download Answers",
            href="/download",
            download="answers.docx",
            id="download-button",
            cls="button",
            style="display: inline-block; padding: 10px 20px; background-color: #4CAF50; color: white; text-decoration: none; border-radius: 4px; border: none; cursor: pointer;",
        ),
        Div(id="download-status"),
    )


def create_hyperlink(
    paragraph: Paragraph, text: str, anchor: str
) -> BaseOxmlElement | etree._Element:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text  # type: ignore
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def remove_temp_file(path: str):
    sleep(5)
    os.unlink(path)


@app.get("/download")
async def download_results(request):
    logger.info(f"Downloading {len(qna['questions'])} answers")
    doc = Document()

    # Add title
    title = doc.add_heading("Answers", level=0)
    title.alignment = 1  # type: ignore

    # Create styles
    styles = doc.styles
    question_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
    question_style.font.bold = True  # type: ignore
    question_style.font.size = Pt(14)  # type: ignore
    answer_style = styles.add_style("Answer", WD_STYLE_TYPE.PARAGRAPH)
    answer_style.font.size = Pt(12)  # type: ignore
    source_style = styles.add_style("Source", WD_STYLE_TYPE.PARAGRAPH)
    source_style.font.size = Pt(10)  # type: ignore
    source_style.font.italic = True  # type: ignore

    # Add questions and answers
    for i, (question, answer, sources) in enumerate(
        zip(qna["questions"], qna["answers"], qna["sources"]), 1
    ):
        doc.add_paragraph(f"Question {i}: {question}", style="Question")
        answer_para = doc.add_paragraph(answer + "\n", style="Answer")
        # Add source references
        for j, source in enumerate(sources, 1):
            create_hyperlink(answer_para, f"[{j}]", f"source_{i}_{j}")
            if j < len(sources):
                answer_para.add_run(", ")
        doc.add_paragraph()  # Add a blank line

    # Add sources section
    doc.add_page_break()
    sources_heading = doc.add_heading("Sources", level=1)
    sources_heading.runs[0].underline = True

    for i, question_sources in enumerate(qna["sources"], 1):
        # Add subheading for each question's sources
        doc.add_heading(f"Question {i}", level=2)

        for j, source in enumerate(question_sources, 1):
            bookmark_name = f"source_{i}_{j}"
            source_para = doc.add_paragraph(style="Source")
            source_run = source_para.add_run(f"{source['name']}")
            source_run.bold = True
            # Add bookmark
            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_start.set(qn("w:id"), f"{i}{j}")
            bookmark_start.set(qn("w:name"), bookmark_name)
            source_para._p.insert(0, bookmark_start)
            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), f"{i}{j}")
            bookmark_end.set(qn("w:name"), bookmark_name)
            source_para._p.append(bookmark_end)
            doc.add_paragraph(source["text"][:SOURCE_TEXT_SIZE], style="Source")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc.save(temp_file.name)
        tmp_path = temp_file.name

    return FileResponse(
        tmp_path,
        filename="answers.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(lambda: remove_temp_file(tmp_path)),
    )


if __name__ == "__main__":
    serve()
