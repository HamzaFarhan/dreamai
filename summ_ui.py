import asyncio
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from fasthtml.common import (
    H3,
    A,
    Button,
    Div,
    FileResponse,
    Form,
    Input,
    Style,
    Titled,
    fast_app,
    serve,
)

app = fast_app(live=True)[0]


# Placeholder function for summary creation
async def create_summary(doc_path):
    # Simulate summary creation with a delay
    await asyncio.sleep(2)
    return f"Summary of {Path(doc_path).name}: This is a placeholder summary. It would contain key points and insights from the document, providing a concise overview of its content."


@app.get("/")
async def show_upload_form():
    return Titled(
        "Company Doc Summarizer",
        Form(
            Input(type="file", name="docs", multiple=True, accept=".pdf,.docx,.txt,.md"),
            Button("Upload and Summarize", type="submit"),
            action="/upload",
            method="post",
            enctype="multipart/form-data",
            hx_post="/upload",
            hx_target="#result",
            hx_indicator="#loading",
        ),
        Div(id="loading", cls="htmx-indicator", _="Summarizing documents..."),
        Div(id="result"),
        Style("""
            form { margin-bottom: 2rem; }
            #loading { display: none; }
            .htmx-request #loading { display: block; }
        """),
    )


@app.post("/upload")
async def handle_doc_upload(request):
    form = await request.form()
    files = form.getlist("docs")

    if not files:
        return Div("No files were uploaded.", cls="error")

    with tempfile.TemporaryDirectory() as temp_dir:
        summaries = []
        for file in files:
            file_path = Path(temp_dir) / file.filename
            with open(file_path, "wb") as f:
                f.write(file.file.read())
            summary = await create_summary(str(file_path))
            summaries.append((file.filename, summary))

        # Create a docx file with summaries
        summary_docx = Path(temp_dir) / "document_summaries.docx"
        document = Document()

        # Add styles
        styles = document.styles
        if "Title" not in styles:
            title_style = styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(18)  # type: ignore
            title_style.font.bold = True  # type: ignore
        else:
            title_style = styles["Title"]

        if "Heading" not in styles:
            heading_style = styles.add_style("Heading", WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(14)  # type: ignore
            heading_style.font.bold = True  # type: ignore
        else:
            heading_style = styles["Heading"]

        # Add title
        document.add_paragraph("Document Summaries", style="Title")

        for filename, summary in summaries:
            # Add filename as heading
            document.add_paragraph(filename, style="Heading")
            # Add summary
            document.add_paragraph(summary)
            # Add a line break between summaries
            document.add_paragraph()

        # Save the document
        document.save(str(summary_docx))

        # Move the file to a location where it can be downloaded
        download_path = Path("static/downloads")
        download_path.mkdir(parents=True, exist_ok=True)
        shutil.move(summary_docx, download_path / "document_summaries.docx")

    return Div(
        H3("Summaries created successfully!"),
        A(
            "Download Summaries",
            href="/static/downloads/document_summaries.docx",
            cls="button",
        ),
        cls="success",
    )


@app.get("/static/downloads/document_summaries.docx")
async def download_summaries(request):
    return FileResponse(
        "static/downloads/document_summaries.docx", filename="document_summaries.docx"
    )


serve()
