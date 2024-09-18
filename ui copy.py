import os
import random
import tempfile
from enum import StrEnum
from pathlib import Path
from time import sleep

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from fasthtml.common import (
    H1,
    H3,
    H5,
    H6,
    A,
    Button,
    Div,
    FileResponse,
    Form,
    Img,
    Input,
    Li,
    Link,
    Main,
    Ol,
    Option,
    P,
    Select,
    Ul,
    fast_app,
    serve,
)
from loguru import logger
from lxml import etree  # type: ignore
from starlette.background import BackgroundTask

from dreamai.md_utils import docs_to_md
from dreamai.settings import ModalSettings

modal_settings = ModalSettings()

LANCE_DIR = Path("/home/hamza/dev/lance")
# LANCE_DIR = Path("/home/hamza/dev/empty_lance")
SOURCE_TEXT_SIZE = 100_000
MODEL = modal_settings.model

logger.info(f"LANCE_DIR: {LANCE_DIR}")
logger.info(f"MODEL: {MODEL}")

table_descriptions = []
questions = []
qna = {}

headers = (Link(rel="stylesheet", href="style_ui.css"),)


app = fast_app(live=True, hdrs=headers)[0]


class Mode(StrEnum):
    RFP = "rfp"
    SECURITY = "security"


def get_sorted_indexes(index_folder: str | Path = LANCE_DIR) -> list[str]:
    index_folder = Path(index_folder)
    if not index_folder.exists():
        return []
    return sorted(
        [dir.name for dir in index_folder.iterdir() if dir.is_dir()],
        key=lambda x: index_folder.joinpath(x).stat().st_mtime,
        reverse=True,
    )


def divider():
    return Div(
        style="border-top: 1px solid #e0e0e0; margin: 30px 0;",
    )


def loader(id: str = "loader"):
    return Img(
        id=id,
        src="SVG-Loaders-master/svg-loaders/ball-triangle.svg",
        cls="ui-indicator",
        style="width: 40px; height: 40px; margin-left: 10px; margin-bottom: 10px;",
    )


def documents_upload(mode: Mode, current_index: str, hx_swap_oob: str | None = None):
    components = [
        H3("Upload Documents (PDF, DOCX, TXT, MD)"),
    ]
    if not current_index:
        components.append(P("No index selected. Please select an index to upload documents."))
    else:
        components += [
            H5(f"Index: {current_index}") if current_index else None,
            Form(
                Input(
                    type="file",
                    name="documents",
                    multiple=True,
                    accept=".pdf,.docx,.txt,.md",
                    id="document-input",
                ),
                loader(),
                hx_post=f"/{mode}/upload-documents",
                hx_trigger="change",
                hx_target="#documents-info",
            ),
            Div(id="documents-info"),
        ]
    return Div(*components, divider(), id="documents-upload", hx_swap_oob=hx_swap_oob)


def questionnaire_upload(mode: Mode, current_index: str, hx_swap_oob: str | None = None):
    components = [H3("Upload Questionnaire CSV")]
    if not current_index:
        components.append(
            P("No index selected. Please select an index to upload questionnaire.")
        )
    else:
        components += [
            H5(f"Index: {current_index}", id="questionnaire-index") if current_index else None,
            Form(
                Input(
                    type="file",
                    name="questionnaire",
                    accept=".csv",
                    id="questionnaire-input",
                ),
                hx_post=f"/{mode}/upload-questionnaire",
                hx_trigger="change",
                hx_target="#questionnaire-info",
            ),
            Div(id="questionnaire-info"),
        ]
    return Div(*components, divider(), id="questionnaire-upload", hx_swap_oob=hx_swap_oob)


def questionnaire_answers(
    mode: Mode = Mode.RFP, questions: list[str] | None = None, hx_swap_oob: str | None = None
):
    if not questions:
        return Div(id="questionnaire-answers", hx_swap_oob=hx_swap_oob)
    return Div(
        Button(
            "Answer Questions",
            hx_get=f"/{mode}/answer-questions",
            hx_indicator="#answer-loader",
            hx_swap="outerHTML",
        ),
        loader(id="answer-loader"),
        id="questionnaire-answers",
        hx_swap_oob=hx_swap_oob,
    )


@app.get("/")
async def root():
    return await home(mode=Mode.RFP)


@app.get("/{mode}")
async def home(mode: Mode = Mode.RFP, index: str = ""):
    indexes = get_sorted_indexes()
    current_index = index or indexes[0] if indexes else index
    container_components = [
        H1(f"{mode.upper() if len(mode) <= 3 else mode.title()} Questionnaire Tool"),
        divider(),
        Div(
            Div(
                Select(
                    *[
                        Option(index, value=index, selected=index == current_index)
                        for index in indexes
                    ],
                    name="index",
                    id="index-select",
                    hx_post=f"/{mode}/select-index",
                    hx_target="#index-info",
                    hx_trigger="change, load",
                    cls="dropdown",
                )
                if indexes
                else None,
                Input(
                    type="text",
                    name="index",
                    placeholder="Create New Index",
                    hx_post=f"/{mode}/select-index",
                    hx_target="#index-info",
                    hx_trigger="keyup[key=='Enter']",
                ),
                id="index-controls",
            ),
            Div(id="index-info"),
        ),
        divider(),
        documents_upload(mode=mode, current_index=current_index),
        questionnaire_upload(mode=mode, current_index=current_index),
        questionnaire_answers(mode=mode),
    ]
    return Main(*container_components, cls="container")


@app.get("/{mode}/index-info")
def index_info(mode: Mode, index: str, index_folder: str | Path = LANCE_DIR):
    index = index.strip().replace(" ", "_")
    existing_tables = get_sorted_indexes(Path(index_folder) / index)
    div = [H5(f"Index: {index}")]
    if existing_tables:
        div.append(P(f"{len(existing_tables)} Existing Data Sources:"))
        div.append(
            Ul(
                *[Li(table) for table in existing_tables],
                style="overflow-y: auto; max-height: 180px;",
            )
        )
    else:
        div.append(P("No data sources found. Upload documents to create data sources."))
    return Div(*div, id="index-info")


@app.post("/{mode}/select-index")
def select_index(mode: Mode = Mode.RFP, index: str = ""):
    return (
        index_info(mode=mode, index=index),
        documents_upload(mode=mode, current_index=index, hx_swap_oob="outerHTML"),
        questionnaire_upload(mode=mode, current_index=index, hx_swap_oob="outerHTML"),
        questionnaire_answers(mode=mode, questions=None, hx_swap_oob="outerHTML"),
    )


@app.post("/{mode}/upload-documents")
async def upload_documents(mode: Mode, request):
    form = await request.form()
    documents = form.getlist("documents")
    logger.info(documents)
    uploaded_files = []
    sleep(1)
    with tempfile.TemporaryDirectory() as temp_dir:
        for file in documents:
            logger.info(file)
            file_name = file.filename
            file_path = Path(temp_dir) / file_name
            with open(file_path, "wb") as f:
                f.write(file.file.read())
            logger.success(f"{file_name} MD:\n\n{docs_to_md(file_path)[0].markdown}")
            uploaded_files.append(file_name)
        logger.info(f"Processed {len(uploaded_files)} files")
    return Div(
        H6(f"Uploaded {len(uploaded_files)} files"),
        Ul(*[Li(file) for file in uploaded_files]),
        hx_get=f"/{mode}/index-info",
        hx_trigger="load",
        hx_target="#index-info",
        hx_include="#index-select",
    )


@app.post("/{mode}/upload-questionnaire")
async def upload_questionnaire(mode: Mode, request):
    global questions
    form = await request.form()
    csv_file = form.get("questionnaire")
    with tempfile.TemporaryDirectory() as temp_dir:
        file_path = Path(temp_dir) / csv_file.filename
        with open(file_path, "wb") as f:
            f.write(csv_file.file.read())

        df = pd.read_csv(file_path)
        questions = df.iloc[:, 0].tolist()
        question_count = len(questions)

        logger.info(f"Uploaded CSV file: {csv_file.filename} with {question_count} questions")
        logger.info(f"Stored {question_count} questions in global variable")

    return Div(
        P(f"Questionnaire file '{csv_file.filename}' uploaded successfully."),
        P(f"Number of questions: {question_count}"),
        Ol(
            *[
                # Li(question[:100] + "..." if len(question) > 100 else question)
                Li(question)
                for question in questions
            ],
            style="overflow-y: auto; max-height: 250px;",
        ),
    ), questionnaire_answers(questions=questions, hx_swap_oob="outerHTML")


async def _answer_questions(questions: list[str]):
    sleep(1)
    global qna
    qna = {"questions": questions, "answers": [], "sources": []}
    for i, question in enumerate(questions, 1):
        logger.info(f"\nAnswering question: {question}")
        qna["answers"].append(f"Answer to question {i}: ...")
        qna["sources"].append(
            [
                {"name": "filename", "text": f"source {s} ..."}
                for s in range(random.randint(1, 3))
            ]
        )


def create_hyperlink(
    paragraph: Paragraph, text: str, anchor: str
) -> BaseOxmlElement | etree._Element:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text  # type: ignore
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


@app.get("/{mode}/answer-questions")
async def answer_questions(mode: Mode):
    await _answer_questions(questions=questions)
    return A("Download Answers", href=f"/{mode}/download", cls="download-button")


@app.get("/{mode}/download")
async def download_file(mode: Mode):
    logger.info(f"Downloading {len(qna['questions'])} questions")
    doc = Document()

    # Add title
    title = doc.add_heading("Questionnaire Answers", level=0)
    title.alignment = 1  # type: ignore

    # Create styles
    styles = doc.styles
    question_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
    question_style.font.bold = True  # type: ignore
    question_style.font.size = Pt(14)  # type: ignore
    answer_style = styles.add_style("Answer", WD_STYLE_TYPE.PARAGRAPH)
    answer_style.font.size = Pt(12)  # type: ignore
    source_style = styles.add_style("Source", WD_STYLE_TYPE.PARAGRAPH)
    source_style.font.size = Pt(10)  # type: ignore
    source_style.font.italic = True  # type: ignore

    # Add questions and answers
    for i, (question, answer, sources) in enumerate(
        zip(qna["questions"], qna["answers"], qna["sources"]), 1
    ):
        doc.add_paragraph(f"Question {i}: {question}", style="Question")
        answer_para = doc.add_paragraph(answer + "\n", style="Answer")
        # Add source references
        for j, source in enumerate(sources, 1):
            create_hyperlink(answer_para, f"[{j}]", f"source_{i}_{j}")
            if j < len(sources):
                answer_para.add_run(", ")
        doc.add_paragraph()  # Add a blank line

    # Add sources section
    doc.add_page_break()
    sources_heading = doc.add_heading("Sources", level=1)
    sources_heading.runs[0].underline = True

    for i, question_sources in enumerate(qna["sources"], 1):
        if not question_sources:
            continue
        # Add subheading for each question's sources
        doc.add_heading(f"Question {i}", level=2)

        for j, source in enumerate(question_sources, 1):
            bookmark_name = f"source_{i}_{j}"
            source_para = doc.add_paragraph(style="Source")
            source_run = source_para.add_run(f"{source['name']}")
            source_run.bold = True
            # Add bookmark
            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_start.set(qn("w:id"), f"{i}{j}")
            bookmark_start.set(qn("w:name"), bookmark_name)
            source_para._p.insert(0, bookmark_start)
            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), f"{i}{j}")
            bookmark_end.set(qn("w:name"), bookmark_name)
            source_para._p.append(bookmark_end)
            doc.add_paragraph(source["text"][:SOURCE_TEXT_SIZE], style="Source")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc.save(temp_file.name)

    return FileResponse(
        temp_file.name,
        filename="answers.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(lambda: os.unlink(temp_file.name)),
    )


if __name__ == "__main__":
    serve()
