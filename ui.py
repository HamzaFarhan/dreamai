import json
import os
import shutil
import tempfile
from enum import StrEnum
from pathlib import Path

import lancedb
import pandas as pd
from burr.core import Application
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from fasthtml.common import (
    H1,
    H3,
    H5,
    H6,
    A,
    Button,
    Div,
    FileResponse,
    Form,
    Img,
    Input,
    Li,
    Link,
    Main,
    Ol,
    Option,
    P,
    Select,
    Ul,
    fast_app,
    serve,
)
from loguru import logger
from lxml import etree  # type: ignore
from starlette.background import BackgroundTask

from dreamai.rag import application
from dreamai.rag_utils import add_data_with_descriptions
from dreamai.settings import ModalSettings
from dreamai.utils import flatten_list

modal_settings = ModalSettings()

LANCE_DIR = modal_settings.lance_dir
# LANCE_DIR = Path("/home/hamza/dev/empty_lance")
SOURCE_TEXT_SIZE = 100_000
MODEL = modal_settings.model

logger.info(f"LANCE_DIR: {LANCE_DIR}")
logger.info(f"MODEL: {MODEL}")

table_descriptions = []
questions = []
qna = {}

headers = (Link(rel="stylesheet", href="style_ui.css"),)


app = fast_app(live=True, hdrs=headers)[0]


class Mode(StrEnum):
    RFP = "rfp"
    SECURITY = "security"


def get_sorted_indexes(index_folder: str | Path = LANCE_DIR) -> list[str]:
    index_folder = Path(index_folder)
    if not index_folder.exists():
        return []
    return sorted(
        [dir.name for dir in index_folder.iterdir() if dir.is_dir()],
        key=lambda x: index_folder.joinpath(x).stat().st_mtime,
        reverse=True,
    )


def divider():
    return Div(
        style="border-top: 1px solid #e0e0e0; margin: 30px 0;",
    )


def loader(id: str = "loader"):
    return Img(
        id=id,
        src="SVG-Loaders-master/svg-loaders/ball-triangle.svg",
        cls="ui-indicator",
        style="width: 40px; height: 40px; margin-left: 10px; margin-bottom: 10px;",
    )


def documents_upload(mode: Mode, current_index: str, hx_swap_oob: str | None = None):
    components = [
        H3("Upload Documents (PDF, DOCX, TXT, MD)"),
    ]
    if not current_index:
        components.append(P("No index selected. Please select an index to upload documents."))
    else:
        components += [
            H5(f"Index: {current_index}") if current_index else None,
            Form(
                Input(
                    type="file",
                    name="documents",
                    multiple=True,
                    accept=".pdf,.docx,.txt,.md",
                    id="document-input",
                ),
                loader(),
                hx_post=f"/{mode}/upload-documents",
                hx_trigger="change",
                hx_target="#documents-info",
            ),
            Div(id="documents-info"),
        ]
    return Div(*components, divider(), id="documents-upload", hx_swap_oob=hx_swap_oob)


def questionnaire_upload(mode: Mode, current_index: str, hx_swap_oob: str | None = None):
    components = [H3("Upload Questionnaire CSV")]
    if not current_index:
        components.append(
            P("No index selected. Please select an index to upload questionnaire.")
        )
    else:
        components += [
            H5(f"Index: {current_index}", id="questionnaire-index") if current_index else None,
            Form(
                Input(
                    type="file",
                    name="questionnaire",
                    accept=".csv",
                    id="questionnaire-input",
                ),
                hx_post=f"/{mode}/upload-questionnaire",
                hx_trigger="change",
                hx_target="#questionnaire-info",
            ),
            Div(id="questionnaire-info"),
        ]
    return Div(*components, divider(), id="questionnaire-upload", hx_swap_oob=hx_swap_oob)


def questionnaire_answers(
    mode: Mode, questions: list[str] | None = None, hx_swap_oob: str | None = None
):
    if not questions:
        return Div(id="questionnaire-answers", hx_swap_oob=hx_swap_oob)
    return Div(
        Button(
            "Answer Questions",
            hx_get=f"/{mode}/answer-questions",
            hx_indicator="#answer-loader",
            hx_swap="outerHTML",
        ),
        loader(id="answer-loader"),
        id="questionnaire-answers",
        hx_swap_oob=hx_swap_oob,
    )


@app.get("/")
async def root():
    return await home(mode=Mode.RFP)


@app.get("/{mode}")
async def home(mode: Mode, index: str = "", hx_swap_oob: str | None = None):
    indexes = get_sorted_indexes()
    current_index = index or indexes[0] if indexes else index
    container_components = [
        H1(f"{mode.upper() if len(mode) <= 3 else mode.title()} Questionnaire Tool"),
        divider(),
        Div(
            Div(
                Select(
                    *[
                        Option(index, value=index, selected=index == current_index)
                        for index in indexes
                    ],
                    name="index",
                    id="index-select",
                    hx_post=f"/{mode}/select-index",
                    hx_target="#index-info",
                    hx_trigger="change, load",
                    cls="dropdown",
                )
                if indexes
                else None,
                Input(
                    type="text",
                    name="index",
                    placeholder="Create New Index",
                    hx_post=f"/{mode}/select-index",
                    hx_target="#index-info",
                    hx_trigger="keyup[key=='Enter']",
                ),
                id="index-controls",
            ),
            Div(id="index-info"),
        ),
        divider(),
        documents_upload(mode=mode, current_index=current_index),
        questionnaire_upload(mode=mode, current_index=current_index),
        questionnaire_answers(mode=mode),
    ]
    return Main(*container_components, cls="container", id="home", hx_swap_oob=hx_swap_oob)


@app.get("/{mode}/index-info")
def index_info(mode: Mode, index: str, index_folder: str | Path = LANCE_DIR):
    global lance_db
    index = index.strip().replace(" ", "_")
    lance_db = lancedb.connect(Path(index_folder) / index)
    existing_tables = list(lance_db.table_names())
    div = [H5(f"Index: {index}")]
    if existing_tables:
        div.append(P(f"{len(existing_tables)} Existing Data Sources:"))
        div.append(
            Ul(
                *[Li(table) for table in existing_tables],
                style="overflow-y: auto; max-height: 180px;",
            )
        )
    else:
        div.append(P("No data sources found. Upload documents to create data sources."))
    div.append(
        A(
            "Delete Index",
            hx_delete=f"/{mode}/delete-index?index={index}",
            hx_confirm=f"Are you sure you want to delete {index}?",
            hx_target="#home",
            style="display: inline-block; padding: 10px 20px; background-color: white; color: red; text-decoration: none; border-radius: 4px; border: none; cursor: pointer; font-weight: bold;",
        )
    )
    return Div(*div, id="index-info")


@app.delete("/{mode}/delete-index")
async def delete_index(mode: Mode, index: str, index_folder: str | Path = LANCE_DIR):
    index_path = Path(index_folder) / index
    if index_path.exists():
        shutil.rmtree(index_path)
    return await home(mode=mode)


@app.post("/{mode}/select-index")
def select_index(mode: Mode, index: str):
    return (
        index_info(mode=mode, index=index),
        documents_upload(mode=mode, current_index=index, hx_swap_oob="outerHTML"),
        questionnaire_upload(mode=mode, current_index=index, hx_swap_oob="outerHTML"),
        questionnaire_answers(mode=mode, questions=None, hx_swap_oob="outerHTML"),
    )


@app.post("/{mode}/upload-documents")
async def upload_documents(mode: Mode, request):
    form = await request.form()
    documents = form.getlist("documents")
    logger.info(documents)
    uploaded_files = []
    with tempfile.TemporaryDirectory() as temp_dir:
        for file in documents:
            # logger.info(file)
            file_name = file.filename
            file_path = Path(temp_dir) / file_name
            with open(file_path, "wb") as f:
                f.write(file.file.read())
            if lance_db.table_names() and file_name in lance_db.table_names():
                uploaded_files.append(f"{file_name} (already exists)")
                continue
            uploaded_files.append(file_name)
            add_data_with_descriptions(
                model=MODEL,
                lance_db=lance_db,
                data=file_path,
                table_descriptions=table_descriptions,
            )
        logger.info(f"Processed {len(uploaded_files)} files")
    index = Path(lance_db.uri).name
    return Div(
        H6(f"Uploaded {len(uploaded_files)} files"),
        Ul(*[Li(file) for file in uploaded_files]),
        hx_get=f"/{mode}/index-info?index={index}",
        hx_trigger="load",
        hx_target="#index-info",
    )


@app.post("/{mode}/upload-questionnaire")
async def upload_questionnaire(mode: Mode, request):
    global questions
    form = await request.form()
    csv_file = form.get("questionnaire")
    with tempfile.TemporaryDirectory() as temp_dir:
        file_path = Path(temp_dir) / csv_file.filename
        with open(file_path, "wb") as f:
            f.write(csv_file.file.read())

        df = pd.read_csv(file_path)
        questions = df.iloc[:, 0].tolist()
        question_count = len(questions)

        logger.info(f"Uploaded CSV file: {csv_file.filename} with {question_count} questions")
        logger.info(f"Stored {question_count} questions in global variable")

    return Div(
        P(f"Questionnaire file '{csv_file.filename}' uploaded successfully."),
        P(f"Number of questions: {question_count}"),
        Ol(
            *[
                # Li(question[:100] + "..." if len(question) > 100 else question)
                Li(question)
                for question in questions
            ],
            style="overflow-y: auto; max-height: 250px;",
        ),
    ), questionnaire_answers(mode=mode, questions=questions, hx_swap_oob="outerHTML")


async def _answer_questions(rag_app: Application, questions: list[str]):
    global qna
    qna = {"questions": questions, "answers": [], "sources": []}
    for query in qna["questions"]:
        inputs = {"query": query}
        logger.info(f"\nProcessing query: {query}")
        while True:
            step_result = rag_app.step(inputs=inputs)
            if step_result is None:
                logger.error("Error: rag_app.step() returned None")
                break
            action, result, _ = step_result
            logger.info(f"\nAction: {action.name}\n")
            logger.success(f"RESULT: {result}\n")
            if action.name == "terminate":
                break
            elif action.name in ["update_chat_history"]:
                qna["answers"].append(result["chat_history"][-1]["content"])
                qna["sources"].append(
                    [
                        {k: v for k, v in json.loads(d).items() if k != "index"}
                        for d in set(
                            [
                                json.dumps(s, sort_keys=True)
                                for s in flatten_list(rag_app.state.get("source_docs", []))
                            ]
                        )
                    ]
                )
                break


def create_hyperlink(
    paragraph: Paragraph, text: str, anchor: str
) -> BaseOxmlElement | etree._Element:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text  # type: ignore
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


@app.get("/{mode}/answer-questions")
async def answer_questions(mode: Mode):
    rag_app = application(db=lance_db, model=MODEL, has_web=False, only_data=True)
    await _answer_questions(rag_app=rag_app, questions=questions)
    return A("Download Answers", href=f"/{mode}/download", cls="download-button")


@app.get("/{mode}/download")
async def download_file(mode: Mode):
    logger.info(f"Downloading {len(qna['questions'])} questions")
    doc = Document()

    # Add title
    title = doc.add_heading(
        f"{mode.upper() if len(mode) <= 3 else mode.title()} Questionnaire Answers", level=0
    )
    title.alignment = 1  # type: ignore

    # Create styles
    styles = doc.styles
    question_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
    question_style.font.bold = True  # type: ignore
    question_style.font.size = Pt(14)  # type: ignore
    answer_style = styles.add_style("Answer", WD_STYLE_TYPE.PARAGRAPH)
    answer_style.font.size = Pt(12)  # type: ignore
    source_style = styles.add_style("Source", WD_STYLE_TYPE.PARAGRAPH)
    source_style.font.size = Pt(10)  # type: ignore
    source_style.font.italic = True  # type: ignore

    # Add questions and answers
    for i, (question, answer, sources) in enumerate(
        zip(qna["questions"], qna["answers"], qna["sources"]), 1
    ):
        doc.add_paragraph(f"Question {i}: {question}", style="Question")
        answer_para = doc.add_paragraph(answer + "\n", style="Answer")
        # Add source references
        for j, source in enumerate(sources, 1):
            create_hyperlink(answer_para, f"[{j}]", f"source_{i}_{j}")
            if j < len(sources):
                answer_para.add_run(", ")
        doc.add_paragraph()  # Add a blank line

    # Add sources section
    doc.add_page_break()
    sources_heading = doc.add_heading("Sources", level=1)
    sources_heading.runs[0].underline = True

    for i, question_sources in enumerate(qna["sources"], 1):
        if not question_sources:
            continue
        # Add subheading for each question's sources
        doc.add_heading(f"Question {i}", level=2)

        for j, source in enumerate(question_sources, 1):
            bookmark_name = f"source_{i}_{j}"
            source_para = doc.add_paragraph(style="Source")
            source_run = source_para.add_run(f"{source['name']}")
            source_run.bold = True
            # Add bookmark
            bookmark_start = OxmlElement("w:bookmarkStart")
            bookmark_start.set(qn("w:id"), f"{i}{j}")
            bookmark_start.set(qn("w:name"), bookmark_name)
            source_para._p.insert(0, bookmark_start)
            bookmark_end = OxmlElement("w:bookmarkEnd")
            bookmark_end.set(qn("w:id"), f"{i}{j}")
            bookmark_end.set(qn("w:name"), bookmark_name)
            source_para._p.append(bookmark_end)
            doc.add_paragraph(source["text"][:SOURCE_TEXT_SIZE], style="Source")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc.save(temp_file.name)

    return FileResponse(
        temp_file.name,
        filename="answers.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(lambda: os.unlink(temp_file.name)),
    )


if __name__ == "__main__":
    serve()
