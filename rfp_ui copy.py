import json
import os
import shutil
import tempfile

import lancedb
import pandas as pd
from docx import Document
from fasthtml.common import (
    H1,
    A,
    Button,
    Container,
    Div,
    Form,
    Input,
    Li,
    P,
    Script,
    Ul,
    fast_app,
    serve,
)
from loguru import logger
from starlette.responses import FileResponse

from dreamai.ai import ModelName
from dreamai.rag import application
from dreamai.rag_utils import add_data_with_descriptions
from dreamai.utils import flatten_list

SOURCE_TEXT_SIZE = 100_000
LANCE_URI = "lance/RFP/"
MODEL = ModelName.GEMINI_FLASH

if os.path.exists(LANCE_URI):
    shutil.rmtree(LANCE_URI)

added_files = []
table_descriptions = []
rfp_questions = []
qna = {}
lance_db = lancedb.connect(uri=LANCE_URI)
app = fast_app(live=True)[0]


@app.get("/")
async def home():
    return Container(
        H1("RFP Question Answering Tool"),
        Div(
            P("Upload your company documents (PDF, DOCX, TXT, or MD files)"),
            Form(
                Input(
                    type="file",
                    name="documents",
                    multiple=True,
                    accept=".pdf,.docx,.txt,.md",
                    id="document-input",
                ),
                Button(
                    "Upload Documents",
                    type="submit",
                    id="upload-button",
                    style="display: none;",
                ),
                hx_post="/upload",
                hx_target="#file-list",
                hx_swap="beforeend",
                enctype="multipart/form-data",
            ),
            Div(id="file-list"),
            id="document-upload-section",
        ),
        Div(
            P("Upload your RFP CSV file"),
            Form(
                Input(type="file", name="rfp", accept=".csv", id="rfp-input"),
                Button(
                    "Upload RFP", type="submit", id="rfp-upload-button", style="display: none;"
                ),
                hx_post="/upload-rfp",
                hx_target="#rfp-info",
                hx_swap="innerHTML",
                enctype="multipart/form-data",
            ),
            Div(id="rfp-info"),
            id="rfp-upload-section",
        ),
        Div(
            Button(
                "Process RFP",
                id="process-button",
                hx_post="/process",
                hx_target="#processing-status",
                style="display: none;",
            ),
            Div(id="processing-status"),
            id="processing-section",
        ),
        Script("""
        document.body.addEventListener('htmx:afterRequest', function(evt) {
            if (evt.detail.successful) {
                if (evt.detail.pathInfo.requestPath === '/upload') {
                    document.getElementById('document-input').value = '';
                    document.getElementById('upload-button').style.display = 'none';
                } else if (evt.detail.pathInfo.requestPath === '/upload-rfp') {
                    document.getElementById('rfp-input').value = '';
                    document.getElementById('rfp-upload-button').style.display = 'none';
                }
                updateProcessButton();
            }
        });

        document.getElementById('document-input').addEventListener('change', function(evt) {
            var uploadButton = document.getElementById('upload-button');
            uploadButton.style.display = this.files.length > 0 ? 'inline-block' : 'none';
        });

        document.getElementById('rfp-input').addEventListener('change', function(evt) {
            var uploadButton = document.getElementById('rfp-upload-button');
            uploadButton.style.display = this.files.length > 0 ? 'inline-block' : 'none';
        });

        function updateProcessButton() {
            var rfpInfoEmpty = document.getElementById('rfp-info').children.length === 0;
            var processButton = document.getElementById('process-button');
            processButton.style.display = !rfpInfoEmpty ? 'inline-block' : 'none';
        }

        // Call this function initially to set the correct state
        updateProcessButton();
        """),
    )


@app.post("/upload")
async def upload(request):
    global table_descriptions
    global added_files
    form = await request.form()
    files = form.getlist("documents")
    if not files:
        return Li("No files were uploaded.", _="on load wait 2s then remove me")
    uploaded_files = []
    with tempfile.TemporaryDirectory() as temp_dir:
        for file in files:
            file_name = file.filename
            logger.info(f"ADDED FILES: {added_files}")
            if file_name in added_files:
                continue
            file_path = os.path.join(temp_dir, file_name)
            with open(file_path, "wb") as f:
                f.write(file.file.read())
            uploaded_files.append(file_name)
            add_data_with_descriptions(
                model=MODEL,
                lance_db=lance_db,
                data=file_path,
                table_descriptions=table_descriptions,
            )
            added_files.append(file_name)
        logger.info(f"Uploaded {len(files)} files to {temp_dir}")

    return Ul(*[Li(file) for file in uploaded_files])


@app.post("/upload-rfp")
async def upload_rfp(request):
    global rfp_questions
    form = await request.form()
    rfp_file = form.get("rfp")
    if not rfp_file:
        return Div("No RFP file was uploaded.", _="on load wait 2s then remove me")

    with tempfile.TemporaryDirectory() as temp_dir:
        file_path = os.path.join(temp_dir, rfp_file.filename)
        with open(file_path, "wb") as f:
            f.write(rfp_file.file.read())

        df = pd.read_csv(file_path)
        rfp_questions = df.iloc[:, 0].tolist()
        question_count = len(rfp_questions)

        logger.info(f"Uploaded RFP file: {rfp_file.filename} with {question_count} questions")
        logger.info(f"Stored {question_count} questions in global variable")

    return Div(
        P(f"RFP file '{rfp_file.filename}' uploaded successfully."),
        P(f"Number of questions: {question_count}"),
        Ul(
            *[
                Li(question[:100] + "..." if len(question) > 100 else question)
                for question in rfp_questions[:5]
            ]
        ),
        P("..." if question_count > 5 else ""),
    )


async def process_questions(app, questions):
    global qna
    qna = {"questions": questions, "answers": [], "sources": []}
    for query in qna["questions"]:
        inputs = {"query": query}
        logger.info(f"\nProcessing query: {query}")
        while True:
            step_result = app.step(inputs=inputs)
            if step_result is None:
                logger.error("Error: app.step() returned None")
                break
            action, result, _ = step_result
            logger.info(f"\nAction: {action.name}\n")
            logger.success(f"RESULT: {result}\n")
            if action.name == "terminate":
                break
            elif action.name in ["update_chat_history"]:
                qna["answers"].append(result["chat_history"][-1]["content"])
                qna["sources"].append(
                    [
                        {k: v for k, v in json.loads(d).items() if k != "index"}
                        for d in set(
                            [
                                json.dumps(s, sort_keys=True)
                                for s in flatten_list(app.state.get("source_docs", []))
                            ]
                        )
                    ]
                )
                break


@app.post("/process")
async def process_rfp(request):
    app = application(db=lance_db, model=MODEL, has_web=False, only_data=True)
    await process_questions(app=app, questions=rfp_questions)
    return Div(
        P("Processing complete!"),
        A(
            "Download Results",
            href="/download",
            download="rfp_answers.docx",
            id="download-button",
            cls="button",
            style="display: inline-block; padding: 10px 20px; background-color: #4CAF50; color: white; text-decoration: none; border-radius: 4px; border: none; cursor: pointer;",
        ),
        Div(id="download-status"),
    )


@app.get("/download")
async def download_results(request):
    logger.info(f"Downloading {len(qna['questions'])} questions")
    doc = Document()
    doc.add_heading("RFP Answers", 0)
    for i, question in enumerate(qna["questions"]):
        doc.add_heading(f"Question {i+1}: {question}", level=1)
        doc.add_paragraph(qna["answers"][i])
        doc.add_heading("Sources:", level=2)
        for source in qna["sources"][i]:
            doc.add_heading(source["name"], level=4)
            doc.add_paragraph(f"... {source['text'][:SOURCE_TEXT_SIZE]} ...")
        doc.add_paragraph()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    return FileResponse(
        tmp_path,
        filename="rfp_answers.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.middleware("http")
async def cleanup_temp_file(request, call_next):
    response = await call_next(request)
    if isinstance(response, FileResponse):
        # await response.close()
        os.unlink(response.path)
    return response


if __name__ == "__main__":
    serve()
